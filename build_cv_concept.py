from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
ASSETS = ROOT / "generated_assets"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)

DOCX_PATH = OUT / "editorial_developer_cv_concept.docx"

INK = "12151A"
MUTED = "626A73"
PAPER = "F6F2E8"
ACID = "B8FF3B"
CORAL = "FF6B5E"
BLUE = "4C78FF"
WHITE = "FFFFFF"
HAIRLINE = "D9D4C8"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_free_page_background(section):
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "7")
        tag.set(qn("w:space"), "20")
        tag.set(qn("w:color"), INK)
        pg_borders.append(tag)
    sect_pr.append(pg_borders)


def set_repeat_font(run, name, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(p, before=0, after=0, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


def set_keep(p, next_=False):
    ppr = p._p.get_or_add_pPr()
    if next_:
        ppr.append(OxmlElement("w:keepNext"))
    ppr.append(OxmlElement("w:keepLines"))


def shade_paragraph(p, fill, color=INK, left=140, right=140, top=80, bottom=80):
    ppr = p._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    borders = OxmlElement("w:pBdr")
    edge = OxmlElement("w:left")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), "18")
    edge.set(qn("w:space"), "8")
    edge.set(qn("w:color"), color)
    borders.append(edge)
    ppr.append(borders)
    ind = p.paragraph_format
    ind.left_indent = Pt(left / 20)
    ind.right_indent = Pt(right / 20)
    ind.space_before = Pt(top / 20)
    ind.space_after = Pt(bottom / 20)


def add_rule(doc, color=INK, size=10, before=4, after=8):
    p = doc.add_paragraph()
    set_para_spacing(p, before, after, 1.0)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)
    return p


def add_label(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, color=INK, after=3):
    p = doc.add_paragraph()
    p.alignment = align
    set_para_spacing(p, 0, after, 1.0)
    r = p.add_run(text.upper())
    set_repeat_font(r, "Consolas", 8.5, color, True)
    r.font.all_caps = True
    set_keep(p, next_=True)
    return p


def add_heading(doc, number, title, color=INK, after=6):
    p = doc.add_paragraph(style="Heading 1")
    set_para_spacing(p, 13, after, 1.0)
    p.paragraph_format.keep_with_next = True
    rn = p.add_run(f"{number} / ")
    set_repeat_font(rn, "Consolas", 10, CORAL, True)
    rt = p.add_run(title.upper())
    set_repeat_font(rt, "Bahnschrift", 16, color, True)
    return p


def add_body(doc, text, after=7, left=0, right=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    set_para_spacing(p, 0, after, 1.22)
    p.paragraph_format.left_indent = Inches(left)
    p.paragraph_format.right_indent = Inches(right)
    r = p.add_run(text)
    set_repeat_font(r, "Calibri", 11, INK)
    return p


def add_mono_note(doc, text, align=WD_ALIGN_PARAGRAPH.RIGHT, color=MUTED, before=0, after=5):
    p = doc.add_paragraph()
    p.alignment = align
    set_para_spacing(p, before, after, 1.0)
    r = p.add_run(text)
    set_repeat_font(r, "Consolas", 8.5, color, False, True)
    return p


def get_font(size, bold=False, mono=False):
    candidates = []
    if mono:
        candidates.extend([
            Path("C:/Windows/Fonts/consola.ttf"),
            Path("C:/Windows/Fonts/lucon.ttf"),
        ])
    elif bold:
        candidates.extend([
            Path("C:/Windows/Fonts/bahnschrift.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
        ])
    else:
        candidates.extend([
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/calibri.ttf"),
        ])
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_loop_graphic(path):
    W, H = 1800, 330
    im = Image.new("RGB", (W, H), "#F6F2E8")
    d = ImageDraw.Draw(im)
    labels = [("IDEA", "#B8FF3B"), ("BUILD", "#FFFFFF"), ("SHIP", "#FF6B5E"), ("LEARN", "#FFFFFF")]
    xs = [70, 500, 930, 1360]
    font = get_font(54, bold=True)
    small = get_font(24, mono=True)
    y = 82
    for i, ((label, fill), x) in enumerate(zip(labels, xs)):
        rounded(d, (x, y, x + 295, y + 145), 30, fill, "#12151A", 6)
        bbox = d.textbbox((0, 0), label, font=font)
        d.text((x + (295 - (bbox[2] - bbox[0])) / 2, y + 40), label, font=font, fill="#12151A")
        if i < len(xs) - 1:
            d.line((x + 305, y + 72, xs[i + 1] - 22, y + 72), fill="#12151A", width=6)
            d.polygon([(xs[i+1]-22, y+72), (xs[i+1]-48, y+57), (xs[i+1]-48, y+87)], fill="#12151A")
    d.arc((1200, 20, 1735, 300), 310, 65, fill="#4C78FF", width=9)
    d.polygon([(1460, 42), (1498, 35), (1481, 72)], fill="#4C78FF")
    d.text((69, 270), "a career path, except it loops on purpose", font=small, fill="#626A73")
    im.save(path, dpi=(300, 300))


def make_project_graphic(path):
    W, H = 1800, 255
    im = Image.new("RGB", (W, H), "#12151A")
    d = ImageDraw.Draw(im)
    big = get_font(42, bold=True)
    small = get_font(22, mono=True)
    items = [
        ("MARKETPLACE", "product + full stack", "#B8FF3B"),
        ("UE5 GAME", "systems + feel", "#FF6B5E"),
        ("AI AGENTS", "B2B + useful automation", "#4C78FF"),
    ]
    for i, (title, sub, col) in enumerate(items):
        x = 65 + i * 580
        d.ellipse((x, 75, x + 34, 109), fill=col)
        d.text((x + 55, 55), title, font=big, fill="#FFFFFF")
        d.text((x + 55, 120), sub, font=small, fill="#C9CDD2")
        if i < 2:
            d.line((x + 485, 92, x + 555, 92), fill="#626A73", width=4)
    im.save(path, dpi=(300, 300))


def make_orbit_graphic(path):
    W, H = 1500, 650
    im = Image.new("RGB", (W, H), "#F6F2E8")
    d = ImageDraw.Draw(im)
    cx, cy = W // 2, H // 2
    d.ellipse((cx - 250, cy - 250, cx + 250, cy + 250), outline="#D9D4C8", width=5)
    d.ellipse((cx - 480, cy - 185, cx + 480, cy + 185), outline="#D9D4C8", width=4)
    rounded(d, (cx - 190, cy - 80, cx + 190, cy + 80), 38, "#12151A")
    center_font = get_font(58, bold=True)
    center_box = d.textbbox((0, 0), "BUILD", font=center_font)
    d.text((cx - (center_box[2]-center_box[0])/2, cy - 38), "BUILD", font=center_font, fill="#FFFFFF")
    font = get_font(38, bold=True)
    mono = get_font(22, mono=True)
    nodes = [
        (145, 155, "USERS", "does this help?", "#B8FF3B"),
        (1125, 130, "SYSTEMS", "will it hold?", "#4C78FF"),
        (1180, 430, "FEEL", "is it alive?", "#FF6B5E"),
        (135, 430, "SHIP", "is it real yet?", "#FFFFFF"),
    ]
    for x, y, title, sub, col in nodes:
        rounded(d, (x, y, x + 245, y + 115), 25, col, "#12151A", 4)
        d.text((x + 23, y + 18), title, font=font, fill="#12151A")
        d.text((x + 23, y + 70), sub, font=mono, fill="#12151A")
    d.text((548, 602), "the stack is technical; the loop is human", font=mono, fill="#626A73")
    im.save(path, dpi=(300, 300))


def make_stack_graphic(path):
    W, H = 1800, 270
    im = Image.new("RGB", (W, H), "#FFFFFF")
    d = ImageDraw.Draw(im)
    labels = [("WEB", ACID), ("MOBILE", "FFFFFF"), ("GAMES", CORAL), ("AI", "FFFFFF"), ("BACKEND", BLUE), ("PRODUCT", "FFFFFF")]
    font = get_font(38, bold=True)
    x = 38
    for label, col in labels:
        bbox = d.textbbox((0, 0), label, font=font)
        w = bbox[2] - bbox[0] + 75
        rounded(d, (x, 72, x + w, 190), 58, f"#{col}", "#12151A", 5)
        d.text((x + 37, 107), label, font=font, fill="#12151A" if col != BLUE else "#FFFFFF")
        x += w + 22
    im.save(path, dpi=(300, 300))


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.70)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.36)
    sec.footer_distance = Inches(0.36)
    sec.different_first_page_header_footer = False
    set_cell_free_page_background(sec)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ("Heading 1", 16, INK, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, INK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Bahnschrift"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Quiet editorial running label: no decorative rule.
    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(hp, 0, 0, 1.0)
    hr = hp.add_run("INDEPENDENT DEVELOPER  /  PROFILE DRAFT")
    set_repeat_font(hr, "Consolas", 7.5, MUTED, True)

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(fp, 0, 0, 1.0)
    fr = fp.add_run("[YOURNAME.DEV]   ·   [EMAIL]   ·   [CITY / REMOTE]")
    set_repeat_font(fr, "Consolas", 7.5, MUTED)


def add_hero(doc, loop_path):
    add_label(doc, "profile / not a chronology", color=CORAL, after=8)

    p = doc.add_paragraph()
    set_para_spacing(p, 0, 3, 0.92)
    set_keep(p, next_=True)
    r1 = p.add_run("I BUILD THINGS\n")
    set_repeat_font(r1, "Bahnschrift", 32, INK, True)
    r2 = p.add_run("UNTIL THEY FEEL ")
    set_repeat_font(r2, "Bahnschrift", 32, INK, True)
    r3 = p.add_run("INEVITABLE.")
    set_repeat_font(r3, "Bahnschrift", 32, BLUE, True)

    p = doc.add_paragraph()
    set_para_spacing(p, 3, 8, 1.12)
    p.paragraph_format.right_indent = Inches(0.55)
    r = p.add_run("Independent full-stack developer / vibecoder / product-minded troublemaker. I make games, web products, mobile apps and the systems behind them.")
    set_repeat_font(r, "Calibri", 13.2, INK, False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 2, 4, 1.0)
    r = p.add_run()
    r.add_picture(str(loop_path), width=Inches(6.65))

    add_mono_note(doc, "// not linear. still moving forward.", after=7)

    p = doc.add_paragraph()
    set_para_spacing(p, 0, 8, 1.18)
    p.paragraph_format.left_indent = Inches(0.48)
    p.paragraph_format.right_indent = Inches(0.48)
    r = p.add_run("“I do the product thinking, the code, the weird middle part, and the last 10% that turns a demo into something people can actually use.”")
    set_repeat_font(r, "Georgia", 12, INK, False, True)
    shade_paragraph(p, ACID, INK, top=70, bottom=70)


def add_project(doc, marker, title, subtitle, text, accent):
    p = doc.add_paragraph()
    set_para_spacing(p, 7, 2, 1.0)
    p.paragraph_format.keep_with_next = True
    rm = p.add_run(marker + "  ")
    set_repeat_font(rm, "Consolas", 9, accent, True)
    rt = p.add_run(title)
    set_repeat_font(rt, "Bahnschrift", 14, INK, True)
    rs = p.add_run("  /  " + subtitle)
    set_repeat_font(rs, "Consolas", 8.5, MUTED, False, True)
    p2 = add_body(doc, text, after=4, left=0.22, right=0.18, align=WD_ALIGN_PARAGRAPH.LEFT)
    p2.paragraph_format.keep_together = True


def build():
    loop_path = ASSETS / "build_loop.png"
    project_path = ASSETS / "project_ribbon.png"
    orbit_path = ASSETS / "build_orbit.png"
    stack_path = ASSETS / "stack_capsules.png"
    make_loop_graphic(loop_path)
    make_project_graphic(project_path)
    make_orbit_graphic(orbit_path)
    make_stack_graphic(stack_path)

    doc = Document()
    configure_document(doc)

    add_hero(doc, loop_path)
    add_heading(doc, "01", "A trail of shipped things", after=5)
    add_body(
        doc,
        "My career is easier to understand as a sequence of things I cared enough to make real. Each project pulled me into a different layer of the stack; each one made the next project more ambitious.",
        after=5,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 2, 3, 1.0)
    p.add_run().add_picture(str(project_path), width=Inches(6.7))

    add_project(
        doc,
        "A",
        "gas.green",
        "marketplace",
        "I helped turn gas.green into a working marketplace, moving between product logic, user flows, frontend and backend until the parts became one usable thing. It is the kind of project I like most: commercial, technical, and real.",
        ACID,
    )
    add_project(
        doc,
        "B",
        "Untitled UE5 project",
        "game",
        "I built a playable world in Unreal Engine 5, working where code, systems, pacing and feel collide. Games make every shortcut visible: if a system is not coherent, the player feels it before they can explain it.",
        CORAL,
    )
    add_project(
        doc,
        "C",
        "AI agents startup",
        "B2B",
        "I worked on B2B AI agents designed around actual company workflows. The goal was not AI as decoration; it was automation that could carry useful work, fit into existing systems, and earn trust.",
        BLUE,
    )
    add_mono_note(doc, "// yes, I named the project chapters. no, I will not apologize.", after=0)

    doc.add_page_break()

    add_label(doc, "how the work happens", color=BLUE, after=5)
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 4, 0.95)
    r1 = p.add_run("FULL STACK,\n")
    set_repeat_font(r1, "Bahnschrift", 28, INK, True)
    r2 = p.add_run("INCLUDING THE HUMAN PART.")
    set_repeat_font(r2, "Bahnschrift", 28, CORAL, True)

    add_body(
        doc,
        "I like owning the distance between a half-formed idea and a thing somebody can use. That means asking product questions, choosing the simplest architecture that survives reality, and caring about the tiny interactions people notice without knowing why.",
        after=4,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 0, 1, 1.0)
    p.add_run().add_picture(str(orbit_path), width=Inches(6.6))
    add_mono_note(doc, "// architecture diagram, but emotionally accurate", after=5)

    add_heading(doc, "02", "The range is the point", after=4)
    add_body(
        doc,
        "Web, mobile, games, AI, backend and product are not separate identities I switch between. They are different materials. I use whichever combination gets the idea out of my head and into the world with the least ceremony and the most useful result.",
        after=4,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 1, 6, 1.0)
    p.add_run().add_picture(str(stack_path), width=Inches(6.7))

    add_heading(doc, "03", "What I am probably good at", after=4)
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 7, 1.16)
    r = p.add_run(
        "Turning ambiguity into a build plan. Learning the missing part while building it. Moving between product detail and system design. Finding the line between ‘clever’ and ‘actually useful’. Shipping before the idea goes stale."
    )
    set_repeat_font(r, "Calibri", 11, INK)
    shade_paragraph(p, "FFFFFF", CORAL, top=80, bottom=80)

    add_heading(doc, "04", "The part a normal CV calls skills", after=4)
    add_body(
        doc,
        "My stack changes with the problem. I care about languages and frameworks, but I care more about choosing tools I can understand, bend, debug and ship. The names belong below; the reason for using them is always the product.",
        after=5,
    )
    add_mono_note(doc, "[ YOUR STACK → TypeScript / React / Node / … / Unreal Engine 5 / … ]", align=WD_ALIGN_PARAGRAPH.LEFT, color=BLUE, after=7)

    add_rule(doc, CORAL, 12, before=3, after=5)
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 2, 1.05)
    r1 = p.add_run("THE NEXT THING")
    set_repeat_font(r1, "Bahnschrift", 16, INK, True)
    r2 = p.add_run("  should be useful, slightly ambitious, and fun to explain at 1 a.m.")
    set_repeat_font(r2, "Georgia", 10.5, MUTED, False, True)
    add_mono_note(doc, "[YOUR NAME]  ·  [EMAIL]  ·  [GITHUB]  ·  [PORTFOLIO]", align=WD_ALIGN_PARAGRAPH.LEFT, color=INK, after=0)

    props = doc.core_properties
    props.title = "Editorial Developer CV Concept"
    props.subject = "Two-page personal CV concept for an independent full-stack developer"
    props.author = ""
    props.keywords = "CV, developer, full stack, games, apps, AI, editorial"

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
